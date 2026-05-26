"""
Gerador de Parecer Técnico de Análise Documental (PT-AD) — Diagnóstika Engenharia.

Lê um arquivo JSON com os dados do caso e gera um documento .docx no padrão Diagnóstika.

Uso:
    python gerar_pt_ad.py inputs/exemplo_apt304.json

Saída:
    saidas/PT-AD_NNN-AAAA_[cliente].docx
"""

import json
import sys
from datetime import datetime
from pathlib import Path

# Garantir UTF-8 no console (Windows cp1252 quebra com emojis e acentos)
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except (AttributeError, OSError):
    pass

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import estilos
import normas
import pendencias

BRANCO = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================
# UTILITÁRIOS BAIXO NÍVEL (XML helpers)
# ============================================================

def shade_cell(cell, hex_fill: str):
    """Aplica cor de fundo a uma célula de tabela."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex="BFBFBF", size="4"):
    """Adiciona bordas finas cinza-claro a uma célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def aplicar_borda_em_tabela(table, color_hex="BFBFBF"):
    """Aplica bordas em todas as células de uma tabela."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color_hex)


def set_run_style(run, *, font=estilos.FONTE_PADRAO, size=None, bold=False,
                  italic=False, color=None):
    """Aplica estilo a um run (segmento de texto)."""
    run.font.name = font
    if size is not None:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ============================================================
# HEADER / FOOTER
# ============================================================

def montar_header_footer(doc, tipo_doc_legivel: str, cliente_curto: str, mes_ano: str):
    """Configura header e footer para todas as páginas (exceto capa)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header das páginas internas
    header = section.header
    header_par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_par.add_run(
        f"Diagnóstika Engenharia · {tipo_doc_legivel} · {cliente_curto} · {mes_ano}"
    )
    set_run_style(run, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)

    # Footer das páginas internas: tabela 3 colunas (esquerda | centro | direita)
    footer = section.footer
    if footer.paragraphs:
        # Limpa parágrafo padrão
        p = footer.paragraphs[0]
        p.clear()

    footer_table = footer.add_table(rows=1, cols=3, width=Cm(17))
    footer_table.autofit = False
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(6), Cm(6), Cm(5)]
    for i, w in enumerate(col_widths):
        footer_table.columns[i].width = w

    # Coluna 1: Razão social
    c1 = footer_table.rows[0].cells[0]
    c1_par = c1.paragraphs[0]
    c1_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = c1_par.add_run("Diagnóstika Engenharia LTDA")
    set_run_style(r1, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)

    # Coluna 2: Cliente | Mês/Ano
    c2 = footer_table.rows[0].cells[1]
    c2_par = c2.paragraphs[0]
    c2_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = c2_par.add_run(f"{cliente_curto} | {mes_ano}")
    set_run_style(r2, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)

    # Coluna 3: Página X de Y (usando campos do Word)
    c3 = footer_table.rows[0].cells[2]
    c3_par = c3.paragraphs[0]
    c3_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3a = c3_par.add_run("Página ")
    set_run_style(r3a, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)
    _add_page_number_field(c3_par)
    r3b = c3_par.add_run(" de ")
    set_run_style(r3b, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)
    _add_total_pages_field(c3_par)


def _add_page_number_field(paragraph):
    """Insere campo {PAGE} no parágrafo."""
    run = paragraph.add_run()
    set_run_style(run, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def _add_total_pages_field(paragraph):
    """Insere campo {NUMPAGES} no parágrafo."""
    run = paragraph.add_run()
    set_run_style(run, size=estilos.TAM_HEADER_FOOTER, color=estilos.CINZA_MEDIO)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "NUMPAGES"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


# ============================================================
# CAPA
# ============================================================

def montar_capa(doc, dados: dict, logo_path: Path):
    """Monta a capa do documento."""
    cliente = dados["cliente"]
    documento = dados["documento"]
    caso = dados["caso"]

    # Logo centralizado
    par_logo = doc.add_paragraph()
    par_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path.exists():
        run_logo = par_logo.add_run()
        run_logo.add_picture(str(logo_path), height=Cm(estilos.CONFIG["altura_logo_capa_cm"]))
    else:
        run_logo = par_logo.add_run("[LOGO DIAGNÓSTIKA]")
        set_run_style(run_logo, size=Pt(16), bold=True, color=estilos.VERDE_PETROLEO)

    doc.add_paragraph()  # espaço

    # Título principal
    tipo_legivel_curto = _tipo_legivel_curto(documento["tipo"], dados)
    par_tit = doc.add_paragraph()
    par_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = par_tit.add_run(tipo_legivel_curto.upper())
    set_run_style(run_tit, size=estilos.TAM_TITULO_CAPA, bold=True,
                  color=estilos.VERDE_PETROLEO)

    # Subtítulo descritivo
    subtitulo = f"{caso['tipo_solicitacao']} — {caso['unidade']}"
    par_sub = doc.add_paragraph()
    par_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = par_sub.add_run(subtitulo)
    set_run_style(run_sub, size=estilos.TAM_SUBTITULO_CAPA, bold=True,
                  color=estilos.VERDE_PETROLEO)

    # Local
    par_loc = doc.add_paragraph()
    par_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_loc = par_loc.add_run(f"{cliente['nome_completo']} · {cliente['cidade_uf']}")
    set_run_style(run_loc, size=estilos.TAM_LOCAL_CAPA, color=estilos.CINZA_MEDIO)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabela de metadados (8 linhas)
    ref_interna = _formatar_ref_interna(documento)
    data_extenso = _data_extenso(documento["data_emissao"])
    status_doc = documento.get("status_capa", "")
    nome_doc_completo = _nome_documento(documento["tipo"]) + (
        f" — {status_doc}" if status_doc else ""
    )

    metadados = [
        ("Documento", nome_doc_completo),
        ("Ref. Interna", ref_interna),
        ("Data de Emissão", data_extenso),
        ("Elaborado por", estilos.EMPRESA["razao_social"].title().replace("Ltda", "LTDA")),
        ("Destinatário", cliente["destinatario"]),
        ("Unidade", caso["unidade"]),
        ("Proprietário", caso.get("proprietario", "—")),
        ("Revisão", documento["revisao"]),
    ]

    tabela = doc.add_table(rows=len(metadados), cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    tabela.columns[0].width = Cm(5)
    tabela.columns[1].width = Cm(12)

    for i, (chave, valor) in enumerate(metadados):
        row = tabela.rows[i]
        # Célula chave
        c_chave = row.cells[0]
        c_chave.width = Cm(5)
        shade_cell(c_chave, estilos.AZUL_CLARO_TABELA)
        c_chave.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_chave = c_chave.paragraphs[0]
        r_chave = p_chave.add_run(chave)
        set_run_style(r_chave, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERDE_PETROLEO)
        # Célula valor
        c_val = row.cells[1]
        c_val.width = Cm(12)
        c_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(str(valor))
        set_run_style(r_val, size=estilos.TAM_TABELA)

    aplicar_borda_em_tabela(tabela)

    # Espaço grande até o rodapé de confidencialidade
    for _ in range(10):
        doc.add_paragraph()

    par_conf = doc.add_paragraph()
    par_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aviso = estilos.AVISO_CONFIDENCIAL.format(
        destinatario_curto=cliente["nome_completo"].upper()
    )
    run_conf = par_conf.add_run(aviso)
    set_run_style(run_conf, size=estilos.TAM_RESSALVA, italic=True,
                  color=estilos.CINZA_MEDIO)

    # Quebra de página depois da capa
    doc.add_page_break()


# ============================================================
# SEÇÕES NUMERADAS
# ============================================================

def titulo_secao(doc, numero: int, texto: str):
    """Adiciona um título de seção (ex: '1. OBJETO E FINALIDADE')."""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(f"{numero}. {texto.upper()}")
    set_run_style(run, size=estilos.TAM_TITULO_SECAO, bold=True,
                  color=estilos.VERDE_PETROLEO)


def paragrafo_corpo(doc, texto: str):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(estilos.CONFIG["espacamento_paragrafos_pt"])
    par.paragraph_format.line_spacing = estilos.CONFIG["espacamento_linhas"]
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(texto)
    set_run_style(run, size=estilos.TAM_CORPO)


def secao_objeto_finalidade(doc, dados: dict):
    titulo_secao(doc, 1, "Objeto e Finalidade")
    paragrafo_corpo(doc, dados["objeto_finalidade"]["texto_objeto"])
    paragrafo_corpo(doc, dados["objeto_finalidade"]["texto_finalidade"])


def secao_documentos_analisados(doc, dados: dict):
    titulo_secao(doc, 2, "Documentos Analisados")

    docs = dados["documentos_analisados"]
    tabela = doc.add_table(rows=len(docs) + 1, cols=3)
    tabela.autofit = False
    tabela.columns[0].width = Cm(2)
    tabela.columns[1].width = Cm(6)
    tabela.columns[2].width = Cm(9)

    # Cabeçalho
    cab = tabela.rows[0]
    for c, txt in zip(cab.cells, ["Ref.", "Documento", "Descrição"]):
        shade_cell(c, estilos.AZUL_CLARO_TABELA)
        r = c.paragraphs[0].add_run(txt)
        set_run_style(r, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERDE_PETROLEO)

    # Linhas
    for i, d in enumerate(docs, 1):
        row = tabela.rows[i]
        for j, key in enumerate(["ref", "nome", "descricao"]):
            cell = row.cells[j]
            r = cell.paragraphs[0].add_run(d[key])
            set_run_style(r, size=estilos.TAM_TABELA)

    aplicar_borda_em_tabela(tabela)


def secao_escopo_construtivo(doc, dados: dict):
    titulo_secao(doc, 3, "Escopo Declarado e Sistema Construtivo")
    caso = dados["caso"]

    linhas = [
        ("Sistema construtivo", caso["sistema_construtivo"] +
         (" — vedações estruturais não admitem rasgos, cortes ou aberturas."
          if caso.get("alvenaria_estrutural") else "")),
        ("Ambiente", caso["ambiente"] +
         (" — área molhada com sistema de impermeabilização."
          if caso["ambiente"].lower() in ("banheiro", "cozinha", "área de serviço") else "")),
        ("Serviços declarados", caso["servicos_declarados"]),
        ("Prazo", f"{caso['prazo_declarado']} (memorial)."),
        ("ART", f"Execução de reforma de edificação — {caso['art_metragem']}."),
    ]

    tabela = doc.add_table(rows=len(linhas), cols=2)
    tabela.autofit = False
    tabela.columns[0].width = Cm(5)
    tabela.columns[1].width = Cm(12)

    for i, (chave, valor) in enumerate(linhas):
        row = tabela.rows[i]
        c_chave = row.cells[0]
        shade_cell(c_chave, estilos.AZUL_CLARO_TABELA)
        c_chave.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r_chave = c_chave.paragraphs[0].add_run(chave)
        set_run_style(r_chave, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERDE_PETROLEO)

        c_val = row.cells[1]
        c_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r_val = c_val.paragraphs[0].add_run(valor)
        set_run_style(r_val, size=estilos.TAM_TABELA)

    aplicar_borda_em_tabela(tabela)


def secao_pendencias(doc, dados: dict):
    titulo_secao(doc, 4, "Pendências Documentais Identificadas")

    lista_pendencias = dados["pendencias"]
    qtd = len(lista_pendencias)

    par_intro = doc.add_paragraph()
    r_intro = par_intro.add_run(
        f"Foram identificadas {qtd:02d} pendências nos documentos recebidos."
    )
    set_run_style(r_intro, size=estilos.TAM_CORPO, bold=True)

    for item in lista_pendencias:
        base = pendencias.mesclar_pendencia(
            item["codigo_base"],
            item.get("override")
        )
        # Resolve placeholders {NBR XXXX} nas strings da pendência
        for chave in ("constatacao", "nao_conformidade", "exigencia"):
            if chave in base:
                base[chave] = pendencias.resolver_citacoes(base[chave], normas)
        _adicionar_bloco_pendencia(doc, item["numero"], base)


def _adicionar_bloco_pendencia(doc, numero: str, conteudo: dict):
    """Cria a tabela 4-linhas de uma pendência (PD-XX)."""
    tabela = doc.add_table(rows=4, cols=2)
    tabela.autofit = False
    tabela.columns[0].width = Cm(3.5)
    tabela.columns[1].width = Cm(13.5)

    # Linha 1: cabeçalho PD-XX | Título
    cab_row = tabela.rows[0]
    c_num = cab_row.cells[0]
    c_tit = cab_row.cells[1]
    shade_cell(c_num, "A04040")  # vermelho-tijolo
    shade_cell(c_tit, "A04040")
    r_num = c_num.paragraphs[0].add_run(numero)
    set_run_style(r_num, size=estilos.TAM_SUBTITULO_SECAO, bold=True, color=BRANCO)
    r_tit = c_tit.paragraphs[0].add_run(conteudo["titulo"])
    set_run_style(r_tit, size=estilos.TAM_SUBTITULO_SECAO, bold=True, color=BRANCO)

    # Linhas 2-4
    for i, (rotulo, chave) in enumerate(
        [("Constatação", "constatacao"),
         ("Não conformidade", "nao_conformidade"),
         ("Exigência", "exigencia")], start=1
    ):
        row = tabela.rows[i]
        c_lab = row.cells[0]
        shade_cell(c_lab, estilos.AZUL_CLARO_TABELA)
        r_lab = c_lab.paragraphs[0].add_run(rotulo)
        set_run_style(r_lab, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERDE_PETROLEO)

        c_txt = row.cells[1]
        r_txt = c_txt.paragraphs[0].add_run(conteudo[chave])
        set_run_style(r_txt, size=estilos.TAM_TABELA)

    aplicar_borda_em_tabela(tabela)

    # Espaço entre pendências
    doc.add_paragraph()


def secao_quadro_resumo(doc, dados: dict):
    titulo_secao(doc, 5, "Quadro Resumo")

    linhas = dados["quadro_resumo"]
    tabela = doc.add_table(rows=len(linhas) + 1, cols=3)
    tabela.autofit = False
    tabela.columns[0].width = Cm(2)
    tabela.columns[1].width = Cm(7)
    tabela.columns[2].width = Cm(8)

    # Cabeçalho
    cab = tabela.rows[0]
    for c, txt in zip(cab.cells, ["Ref.", "Tema", "Pendência"]):
        shade_cell(c, estilos.AZUL_CLARO_TABELA)
        r = c.paragraphs[0].add_run(txt)
        set_run_style(r, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERDE_PETROLEO)

    # Linhas
    for i, item in enumerate(linhas, 1):
        row = tabela.rows[i]
        for j, key in enumerate(["ref", "tema", "pendencia"]):
            cell = row.cells[j]
            r = cell.paragraphs[0].add_run(item[key])
            set_run_style(r, size=estilos.TAM_TABELA,
                          bold=(j == 0))

    aplicar_borda_em_tabela(tabela)


def secao_conclusao(doc, dados: dict):
    titulo_secao(doc, 6, "Conclusão e Encaminhamentos")

    conc = dados["conclusao"]

    # Texto principal da conclusão (com status em negrito)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto = conc["texto_conclusao"]
    # Destacar status no texto
    status = conc["status"]
    if status in texto:
        antes, depois = texto.split(status, 1)
        r1 = par.add_run(antes)
        set_run_style(r1, size=estilos.TAM_CORPO, bold=True)
        r2 = par.add_run(status)
        set_run_style(r2, size=estilos.TAM_CORPO, bold=True,
                      color=estilos.VERMELHO_TIJOLO)
        r3 = par.add_run(depois)
        set_run_style(r3, size=estilos.TAM_CORPO, bold=True)
    else:
        r = par.add_run(texto)
        set_run_style(r, size=estilos.TAM_CORPO, bold=True)

    # Ressalva
    if conc.get("texto_ressalva"):
        par_res = doc.add_paragraph()
        par_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_res = par_res.add_run(conc["texto_ressalva"])
        set_run_style(r_res, size=estilos.TAM_CORPO)

    # Caixa de encaminhamentos
    doc.add_paragraph()
    par_titEnc = doc.add_paragraph()
    r_titEnc = par_titEnc.add_run("ENCAMINHAMENTOS")
    set_run_style(r_titEnc, size=estilos.TAM_CORPO, bold=True,
                  color=estilos.VERDE_PETROLEO)

    for enc in conc["encaminhamentos"]:
        par_enc = doc.add_paragraph()
        par_enc.paragraph_format.left_indent = Cm(0.5)
        r_enc = par_enc.add_run(f"■ {enc}")
        set_run_style(r_enc, size=estilos.TAM_CORPO)


def secao_responsabilidade(doc, dados: dict):
    titulo_secao(doc, 7, "Responsabilidade Técnica")

    par_rs = doc.add_paragraph()
    r_rs = par_rs.add_run(estilos.EMPRESA["razao_social"])
    set_run_style(r_rs, size=estilos.TAM_CORPO, bold=True,
                  color=estilos.VERDE_PETROLEO)

    par_cnpj = doc.add_paragraph()
    r_cnpj = par_cnpj.add_run(
        f"CNPJ {estilos.EMPRESA['cnpj']} | {estilos.EMPRESA['cidade_uf']}"
    )
    set_run_style(r_cnpj, size=estilos.TAM_CORPO)

    doc.add_paragraph()
    par_resp = doc.add_paragraph()
    r_resp1 = par_resp.add_run("Responsável pela elaboração:")
    set_run_style(r_resp1, size=estilos.TAM_CORPO, bold=True)
    par_resp2 = doc.add_paragraph()
    r_resp2 = par_resp2.add_run(estilos.FUNCAO_TECNICA_PT_AD)
    set_run_style(r_resp2, size=estilos.TAM_CORPO)

    doc.add_paragraph()
    par_data = doc.add_paragraph()
    data_extenso = _data_extenso(dados["documento"]["data_emissao"])
    r_data = par_data.add_run(f"{estilos.EMPRESA['cidade']}, {data_extenso}.")
    set_run_style(r_data, size=estilos.TAM_CORPO)

    doc.add_paragraph()
    par_ress = doc.add_paragraph()
    par_ress.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_ress = par_ress.add_run(estilos.RESSALVA_PT_AD)
    set_run_style(r_ress, size=estilos.TAM_RESSALVA, italic=True,
                  color=estilos.CINZA_MEDIO)


# ============================================================
# UTILITÁRIOS DE FORMATAÇÃO DE DADOS
# ============================================================

def _formatar_ref_interna(documento: dict) -> str:
    tipo = documento["tipo"]  # ex: "PT-AD"
    seq = documento["ref_sequencial"]
    ano = documento["ano"]
    return f"DK-{tipo}-{seq:03d}/{ano}"


def _data_extenso(iso_str: str) -> str:
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    d = datetime.fromisoformat(iso_str)
    return f"{d.day} de {meses[d.month - 1]} de {d.year}"


def _mes_ano(iso_str: str) -> str:
    meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
             "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    d = datetime.fromisoformat(iso_str)
    return f"{meses[d.month - 1]}/{d.year}"


def _nome_documento(tipo: str) -> str:
    nomes = {
        "PT-AD": "Parecer Técnico de Análise Documental",
        "PT-NC": "Parecer Técnico de Não Conformidades",
        "LT": "Laudo Técnico",
        "RT": "Relatório Técnico",
    }
    return nomes.get(tipo, f"Documento {tipo}")


def _tipo_legivel_curto(tipo: str, dados: dict) -> str:
    return _nome_documento(tipo)


# ============================================================
# MAIN
# ============================================================

def verificar_normas(dados: dict) -> list:
    """Verifica status das normas usadas. Retorna lista de alertas."""
    codigos = pendencias.coletar_normas_usadas(dados.get("pendencias", []))
    alertas = normas.alertas_para_pareceres(codigos)
    return alertas


def imprimir_alertas_console(alertas: list):
    """Imprime alertas no console (usado durante a geração)."""
    if not alertas:
        print("[OK] Todas as normas citadas estao com status 'vigente_confirmada'.")
        return
    print("=" * 70)
    print("[!!] ALERTAS DE NORMAS - revisar antes de enviar o documento:")
    print("=" * 70)
    for a in alertas:
        print(f"  - {a['codigo']:<25} [{a['severidade']}] {a['status']}")
        if a.get("acao_requerida"):
            print(f"      Acao: {a['acao_requerida']}")
        if a.get("observacao"):
            print(f"      Obs:  {a['observacao']}")
    print("=" * 70)


def adicionar_caixa_alerta_normas(doc, alertas: list):
    """Insere uma caixa de alerta visível no início do documento (após a capa)."""
    if not alertas:
        return

    # Título do alerta
    par = doc.add_paragraph()
    run = par.add_run("⚠ ALERTA — REVISAR CITAÇÕES DE NORMAS ANTES DA EMISSÃO FINAL")
    set_run_style(run, size=estilos.TAM_CORPO, bold=True,
                  color=estilos.VERMELHO_TIJOLO)

    # Tabela com alertas
    tabela = doc.add_table(rows=len(alertas) + 1, cols=3)
    tabela.autofit = False
    tabela.columns[0].width = Cm(3)
    tabela.columns[1].width = Cm(3)
    tabela.columns[2].width = Cm(11)

    cab = tabela.rows[0]
    for c, txt in zip(cab.cells, ["Norma", "Severidade", "Ação requerida"]):
        shade_cell(c, "F8D8D8")
        r = c.paragraphs[0].add_run(txt)
        set_run_style(r, size=estilos.TAM_TABELA, bold=True,
                      color=estilos.VERMELHO_TIJOLO)

    for i, a in enumerate(alertas, 1):
        row = tabela.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(a["codigo"])
        set_run_style(r0, size=estilos.TAM_TABELA, bold=True)
        r1 = row.cells[1].paragraphs[0].add_run(a["severidade"])
        set_run_style(r1, size=estilos.TAM_TABELA, bold=True)
        msg = a.get("acao_requerida") or a.get("mensagem", "")
        r2 = row.cells[2].paragraphs[0].add_run(msg)
        set_run_style(r2, size=estilos.TAM_TABELA)

    aplicar_borda_em_tabela(tabela, color_hex="E0A0A0")

    par_obs = doc.add_paragraph()
    r_obs = par_obs.add_run(
        "Este parecer foi gerado automaticamente. Antes de emissão oficial, "
        "verifique no catálogo ABNT (https://www.abntcatalogo.com.br) a versão "
        "vigente de cada norma marcada acima e atualize o normas_catalogo.json."
    )
    set_run_style(r_obs, size=estilos.TAM_RESSALVA, italic=True,
                  color=estilos.CINZA_MEDIO)

    doc.add_paragraph()


def gerar(input_json_path: Path, output_dir: Path, assets_dir: Path) -> Path:
    with open(input_json_path, "r", encoding="utf-8") as f:
        dados = json.load(f)

    # Verificação de normas ANTES de gerar
    alertas_normas = verificar_normas(dados)
    imprimir_alertas_console(alertas_normas)

    doc = Document()

    # Margens
    section = doc.sections[0]
    section.top_margin = estilos.MARGEM_SUPERIOR
    section.bottom_margin = estilos.MARGEM_INFERIOR
    section.left_margin = estilos.MARGEM_ESQUERDA
    section.right_margin = estilos.MARGEM_DIREITA

    # Header/footer
    tipo_legivel = _nome_documento(dados["documento"]["tipo"])
    mes_ano = _mes_ano(dados["documento"]["data_emissao"])
    montar_header_footer(
        doc,
        tipo_doc_legivel=tipo_legivel,
        cliente_curto=dados["cliente"]["nome_curto"],
        mes_ano=mes_ano,
    )

    # Capa
    logo_path = assets_dir / "logo.png"
    montar_capa(doc, dados, logo_path)

    # Caixa de alerta de normas (logo após a capa, antes das seções)
    if alertas_normas:
        adicionar_caixa_alerta_normas(doc, alertas_normas)

    # Seções
    secao_objeto_finalidade(doc, dados)
    secao_documentos_analisados(doc, dados)
    secao_escopo_construtivo(doc, dados)
    secao_pendencias(doc, dados)
    secao_quadro_resumo(doc, dados)
    secao_conclusao(doc, dados)
    secao_responsabilidade(doc, dados)

    # Salvar
    ref = _formatar_ref_interna(dados["documento"]).replace("/", "-")
    cliente_arq = dados["cliente"]["nome_curto"].replace(" ", "_")
    nome_arq = f"{ref}_{cliente_arq}.docx"
    output_path = output_dir / nome_arq
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Uso: python gerar_pt_ad.py <input.json>")
        sys.exit(1)

    base_dir = Path(__file__).parent
    input_path = Path(sys.argv[1])
    if not input_path.is_absolute():
        input_path = base_dir / input_path

    output_dir = base_dir / "saidas"
    assets_dir = base_dir / "assets"

    print(f"Lendo: {input_path}")
    resultado = gerar(input_path, output_dir, assets_dir)
    print(f"[OK] Documento gerado: {resultado}")
    print(f"     Tamanho: {resultado.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
